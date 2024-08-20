from docx import Document

# Create a new Document
doc = Document()

# Add the main heading with the name
doc.add_heading('Michael Murphy', level=1)

# Add contact information
doc.add_paragraph('Email: mpmurphy2014@gmail.com')
doc.add_paragraph('Phone: (214) 707-9570')
doc.add_paragraph('Location: Dallas, TX')

# Add Professional Summary
doc.add_heading('Professional Summary', level=2)
doc.add_paragraph(
    'Results-oriented Product Owner and Systems Engineer with a proven track record of delivering complex engineering projects in the aerospace and defense sectors. Expertise in DFARS and CMMC compliance, agile methodologies, and DevSecOps practices. Demonstrated ability to satisfy customer requirements while meeting deadlines and budgets, effectively managing stakeholder collaboration and vendor relationships.'
)

# Add Education
doc.add_heading('Education', level=2)
doc.add_paragraph('Master of Arts in Design & Innovation')
doc.add_paragraph('Bachelor of Science in Mechanical Engineering')
doc.add_paragraph('Southern Methodist University, Dallas, TX')
doc.add_paragraph('Graduated: May 2019')

# Add Professional Experience
doc.add_heading('Professional Experience', level=2)

# Experience 1
doc.add_heading('Product Owner, Sr. Systems Engineer', level=3)
doc.add_paragraph('Lockheed Martin, Fort Worth, TX')
doc.add_paragraph('Nov 2021 – Present')
doc.add_paragraph(
    '- Spearheaded the development of an enterprise synchronization and integration pipeline, achieving a 20% increase in project delivery speed by implementing agile frameworks.'
)
doc.add_paragraph(
    '- Successfully coordinated customer collaboration across internal and external teams, resulting in a 15% cost reduction by optimizing procurement strategies.'
)
doc.add_paragraph(
    '- Delivered IT transformation projects on time and within budget by providing standardized frameworks and automation libraries, enhancing customer satisfaction.'
)
doc.add_paragraph(
    '- Implemented Agile and DevSecOps practices, ensuring compliance with DO-178C standards and reducing security vulnerabilities by 30%.'
)
doc.add_paragraph(
    '- Led automation of enterprise PLM tools, achieving an 80% reduction in delivery and update times, which significantly improved operational efficiency.'
)

# Experience 2
doc.add_heading('Innovation Interaction Designer, UX Lead', level=3)
doc.add_paragraph('Lockheed Martin, Fort Worth, TX')
doc.add_paragraph('Jun 2019 – Feb 2022')
doc.add_paragraph(
    '- Developed a DevSecOps accelerator program, increasing organizational maturity by 25% and upskilling teams through targeted training and resources.'
)
doc.add_paragraph(
    '- Designed and implemented an enterprise FOSS tool, improving collaboration and reducing risk, which led to a 40% increase in workflow efficiency.'
)
doc.add_paragraph(
    '- Directed UX initiatives, enhancing IT experiences and replacing legacy systems, resulting in a 30% improvement in user satisfaction.'
)

# Experience 3
doc.add_heading('Mechanical Engineering/Design, MEP', level=3)
doc.add_paragraph('Purdy-McGuire, Inc., Dallas, TX')
doc.add_paragraph('Nov 2016 – Jun 2019')
doc.add_paragraph(
    '- Created detailed engineering floor plans, supporting REVIT Building Information Modeling projects and contributing to a 20% reduction in design errors.'
)

# Add Skills
doc.add_heading('Skills', level=2)
doc.add_paragraph('- Business Development and Strategic Planning')
doc.add_paragraph('- Agile Project Management and DevSecOps')
doc.add_paragraph('- Compliance with DFARS, CMMC, and DO-178C Standards')
doc.add_paragraph('- Budget Management and Forecasting')
doc.add_paragraph('- Stakeholder Collaboration and Vendor Evaluation')
doc.add_paragraph('- UX/UI Design and Digital Transformation')
doc.add_paragraph('- AWS Solutions Architect, CAD, Adobe Creative Suite')

# Add Certifications
doc.add_heading('Certifications', level=2)
doc.add_paragraph('- AWS Solutions Architect Associate')
doc.add_paragraph('- Engineer in Training (EIT)')
doc.add_paragraph('- Top-Secret Clearance')

# Add Additional Information
doc.add_heading('Additional Information', level=2)
doc.add_paragraph(
    '- Proven ability to manage projects with constant customer collaboration, securing funding from executive leadership, and delivering results that exceed expectations.'
)
doc.add_paragraph(
    '- Committed to advancing aerospace technology and contributing to innovative air transportation solutions.'
)

# Save the document
doc.save('Michael_Murphy_Resume.docx')
